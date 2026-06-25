from __future__ import annotations

import re
from dataclasses import dataclass
from typing import Dict, List, Tuple

import pandas as pd

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None

DIMENSIONS: Dict[str, List[str]] = {
    "Digital Scam Infrastructure": [
        "telegram", "facebook", "whatsapp", "wechat", "social media", "online", "website", "link", "group", "platform",
    ],
    "Victim Vulnerability": [
        "victim", "mangsa", "retiree", "teacher", "student", "professional", "greed", "quick", "rich", "profit", "returns", "dividend",
    ],
    "Scammer Operational Mechanism": [
        "scammer", "syndicate", "mastermind", "agent", "call center", "script", "testimonial", "seminar", "package", "slot", "promise",
    ],
    "Financial Laundering Network": [
        "mule", "account", "bank", "atm", "transaction", "transfer", "crypto", "cryptocurrency", "money", "payment", "deposit",
    ],
    "Institutional Response": [
        "police", "pdrm", "bukit aman", "ipk", "ipd", "bnm", "bank negara", "ssm", "skmm", "sc", "nsrc", "dpp", "court", "section 420",
    ],
    "Prevention and Public Awareness": [
        "awareness", "prevention", "campaign", "lecture", "facebook", "tiktok", "pamphlet", "education", "hotline", "mule check", "alert",
    ],
}

WARNING_SIGNS: Dict[str, List[str]] = {
    "Unrealistic high return": ["high return", "profit", "returns", "30%", "300", "400", "short time", "within", "guarantee"],
    "Urgency pressure": ["immediate", "limited", "short period", "act immediately", "quick", "fast", "now"],
    "Telegram/social media recruitment": ["telegram", "facebook", "whatsapp", "social media", "group", "link"],
    "Mule account / third-party account": ["mule", "atm card", "account", "third party", "rent", "borrow", "sell"],
    "Fake testimonial / proof of profit": ["testimonial", "receipt", "proof", "dividend", "profit", "got it"],
    "Unlicensed investment": ["license", "unlicensed", "bank negara", "bnm", "securities commission", "sc", "alert list"],
    "No real business model": ["non-existent", "does not exist", "no business", "no product", "money game", "ponzi", "pyramid"],
}

STAKEHOLDER_ACTIONS = [
    ("PDRM / CCID", "Evidence gathering, investigation, suspect tracing, prosecution support", "High"),
    ("BNM / Banks", "Real-time transaction monitoring, account freezing, mule account detection", "High"),
    ("SSM", "Company registration screening, blacklist matching, legitimacy verification", "Medium-High"),
    ("SKMM / Telcos", "SIM registration control, suspicious number tracing, platform coordination", "Medium-High"),
    ("SC Malaysia", "Investment legitimacy checks, investor alerts, securities enforcement", "Medium-High"),
    ("NSRC", "Rapid scam response, coordination with banks, early blocking of losses", "High"),
    ("Digital Platforms", "Remove scam groups/pages, detect fake testimonials, flag risky offers", "High"),
    ("Public / Victims", "Due diligence, verify license, avoid mule accounts, report early", "High"),
]


def read_text_from_file(file_obj) -> str:
    """Read txt/docx uploads from Streamlit UploadedFile or local path."""
    name = getattr(file_obj, "name", str(file_obj)).lower()
    if name.endswith(".txt"):
        data = file_obj.read() if hasattr(file_obj, "read") else open(file_obj, "rb").read()
        return data.decode("utf-8", errors="ignore")
    if name.endswith(".docx"):
        if Document is None:
            raise RuntimeError("python-docx is required to read DOCX files")
        doc = Document(file_obj)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    raise ValueError("Supported formats: .docx and .txt")


def clean_text(text: str) -> str:
    text = re.sub(r"\s+", " ", text or " ").strip()
    return text


def count_keywords(text: str, keywords: List[str]) -> Tuple[int, Dict[str, int]]:
    lower = text.lower()
    detail = {}
    total = 0
    for kw in keywords:
        n = len(re.findall(r"\b" + re.escape(kw.lower()) + r"\b", lower)) if " " not in kw else lower.count(kw.lower())
        detail[kw] = n
        total += n
    return total, detail


def analyze_texts(texts: Dict[str, str]) -> Dict[str, pd.DataFrame]:
    records = []
    warning_records = []
    quote_records = []
    for doc_name, raw in texts.items():
        text = clean_text(raw)
        sentences = re.split(r"(?<=[.!?])\s+", text)
        for dim, kws in DIMENSIONS.items():
            score, detail = count_keywords(text, kws)
            records.append({"Transcript": doc_name, "Dimension": dim, "Evidence_Count": score})
            # quotation snippets
            hits = []
            for s in sentences:
                if any(kw.lower() in s.lower() for kw in kws):
                    hits.append(s[:260])
                if len(hits) >= 2:
                    break
            quote_records.append({"Transcript": doc_name, "Dimension": dim, "Indicative_Evidence": " | ".join(hits)})
        for sign, kws in WARNING_SIGNS.items():
            score, _ = count_keywords(text, kws)
            warning_records.append({"Transcript": doc_name, "Warning_Sign": sign, "Signal_Count": score})
    dim_df = pd.DataFrame(records)
    warn_df = pd.DataFrame(warning_records)
    quotes_df = pd.DataFrame(quote_records)
    summary = dim_df.groupby("Dimension", as_index=False)["Evidence_Count"].sum().sort_values("Evidence_Count", ascending=False)
    risk = warn_df.groupby("Transcript", as_index=False)["Signal_Count"].sum()
    max_score = max(risk["Signal_Count"].max(), 1)
    risk["Risk_Index_0_100"] = (risk["Signal_Count"] / max_score * 100).round(1)
    return {"dimensions": dim_df, "dimension_summary": summary, "warnings": warn_df, "risk": risk, "quotes": quotes_df}


def build_interaction_edges() -> pd.DataFrame:
    edges = [
        ("Digital Platforms", "Scammer Recruitment", 5),
        ("Scammer Recruitment", "Psychological Manipulation", 5),
        ("Psychological Manipulation", "Victim Decision", 4),
        ("Victim Decision", "Financial Transfer", 5),
        ("Financial Transfer", "Mule Account Network", 5),
        ("Mule Account Network", "Layered Transactions", 4),
        ("Layered Transactions", "Mastermind Concealment", 4),
        ("PDRM / CCID", "Investigation", 5),
        ("Banks / BNM", "Transaction Blocking", 5),
        ("SSM", "Company Legitimacy Screening", 3),
        ("SKMM / Telcos", "SIM and Number Tracing", 3),
        ("SC Malaysia", "Investor Alert", 3),
        ("NSRC", "Rapid Response", 5),
        ("Public Awareness", "Victim Decision", 3),
        ("Digital Platforms", "Scam Content Removal", 4),
    ]
    return pd.DataFrame(edges, columns=["Source", "Target", "Weight"])


def stakeholder_matrix() -> pd.DataFrame:
    return pd.DataFrame(STAKEHOLDER_ACTIONS, columns=["Stakeholder", "Prevention Role", "Priority"])
